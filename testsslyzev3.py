from datetime import datetime, UTC
from pathlib import Path
from typing import List
from colorama import init, Fore, Style
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from sslyze import (
    Scanner,
    ServerScanRequest,
    ScanCommand,
    SslyzeOutputAsJson,
    ServerNetworkLocation,
    ScanCommandAttemptStatusEnum,
    ServerScanStatusEnum,
    ServerScanResult,
    ServerScanResultAsJson,
)
from sslyze.errors import ServerHostnameCouldNotBeResolved
from sslyze.scanner.scan_command_attempt import ScanCommandAttempt
import warnings
from cryptography.utils import CryptographyDeprecationWarning

# Suppress cryptography deprecation warnings
warnings.filterwarnings("ignore", category=CryptographyDeprecationWarning)

# Initialize colorama for colored terminal output
init(autoreset=True)

# Define Mozilla's recommended cipher suites
MOZILLA_SAFE_CIPHERS = {
    "TLS_AES_256_GCM_SHA384",
    "TLS_AES_128_GCM_SHA256",
    "TLS_CHACHA20_POLY1305_SHA256",
    "TLS_ECDHE_ECDSA_WITH_AES_128_GCM_SHA256",
    "TLS_ECDHE_RSA_WITH_AES_128_GCM_SHA256",
    "TLS_ECDHE_ECDSA_WITH_AES_256_GCM_SHA384",
    "TLS_ECDHE_RSA_WITH_AES_256_GCM_SHA384",
}

def print_ascii_banner():
    banner = r"""
    ==========================================
      ____           _         _   ____  _    
     / ___|   _ __  | |   ___ | |_| __ )| |   
     \___ \  | '_ \ | |  / _ \| __|  _ \| |   
      ___) | | |_) || |_|  __/| |_| |_) | |___
     |____/  | .__/ |____\___| \__|____/|_____|
             |_|                              
    Custom SSL Scan by BlkPh0x
    ==========================================
    """
    print(Fore.CYAN + banner + Style.RESET_ALL)

def _print_failed_scan_command_attempt(scan_command_attempt: ScanCommandAttempt) -> None:
    print(
        f"\nError when running scan command: {scan_command_attempt.error_reason}:\n"
        f"{scan_command_attempt.error_trace}"
    )

def mark_and_sort_cipher_suites(cipher_suites):
    safe_suites = []
    unsafe_suites = []

    for suite in cipher_suites:
        cipher_name = suite.cipher_suite.name
        if cipher_name in MOZILLA_SAFE_CIPHERS:
            safe_suites.append(cipher_name)
        else:
            unsafe_suites.append(cipher_name)

    for suite in sorted(safe_suites):
        print(f"{Fore.GREEN}* {suite} - SAFE{Style.RESET_ALL}")

    for suite in sorted(unsafe_suites):
        print(f"{Fore.RED}* {suite} - UNSAFE{Style.RESET_ALL}")

    return sorted(safe_suites), sorted(unsafe_suites)

def write_results_to_docx(results: dict, output_file: str):
    doc = Document()
    doc.add_heading(f"SSL/TLS Scan Results for {results['hostname']}", level=1)

    def add_table(header, rows):
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Test"
        hdr_cells[1].text = "Result"
        for row in rows:
            row_cells = table.add_row().cells
            row_cells[0].text, row_cells[1].text = row
        return table

    for tls_version, ciphers in results["ciphers"].items():
        doc.add_heading(f"{tls_version} Cipher Suites", level=2)
        rows = [(cipher, status) for cipher, status in ciphers]
        add_table(f"{tls_version} Cipher Suites", rows)

    if results.get("additional_checks"):
        doc.add_heading("Additional Checks", level=2)
        rows = results["additional_checks"]
        add_table("Additional Checks", rows)

    doc.save(output_file)
    print(f"\nResults saved to {output_file}")

def main() -> None:
    print_ascii_banner()
    print("=> Starting the scans")
    date_scans_started = datetime.now(UTC)

    servers_input = input("Enter server hostnames separated by commas: ")
    server_list = [host.strip() for host in servers_input.split(",") if host.strip()]

    if not server_list:
        print("No servers provided. Exiting.")
        return

    all_scan_requests = []
    for server in server_list:
        try:
            all_scan_requests.append(ServerScanRequest(
                server_location=ServerNetworkLocation(hostname=server),
                scan_commands=[
                    ScanCommand.SSL_2_0_CIPHER_SUITES,
                    ScanCommand.SSL_3_0_CIPHER_SUITES,
                    ScanCommand.TLS_1_0_CIPHER_SUITES,
                    ScanCommand.TLS_1_1_CIPHER_SUITES,
                    ScanCommand.TLS_1_2_CIPHER_SUITES,
                    ScanCommand.TLS_1_3_CIPHER_SUITES,
                    ScanCommand.CERTIFICATE_INFO,
                    ScanCommand.HEARTBLEED,
                    ScanCommand.SESSION_RESUMPTION,
                    ScanCommand.TLS_COMPRESSION,
                    ScanCommand.TLS_1_3_EARLY_DATA,
                    ScanCommand.OPENSSL_CCS_INJECTION,
                    ScanCommand.TLS_FALLBACK_SCSV,
                    ScanCommand.ROBOT,
                    ScanCommand.SESSION_RENEGOTIATION,
                    ScanCommand.ELLIPTIC_CURVES,
                    ScanCommand.HTTP_HEADERS,
                ]
            ))
        except ServerHostnameCouldNotBeResolved:
            print(f"Error: Could not resolve hostname '{server}'")
            return

    scanner = Scanner()
    scanner.queue_scans(all_scan_requests)

    for server_scan_result in scanner.get_results():
        hostname = server_scan_result.server_location.hostname
        print(f"\n\n****Results for {hostname}****")

        if server_scan_result.scan_status == ServerScanStatusEnum.ERROR_NO_CONNECTIVITY:
            print(f"Error: Could not connect to {hostname}: {server_scan_result.connectivity_error_trace}")
            continue

        assert server_scan_result.scan_result

        result_entry = {"hostname": hostname, "ciphers": {}, "additional_checks": []}

        for version, attempt in [
            ("TLS 1.3", server_scan_result.scan_result.tls_1_3_cipher_suites),
            ("TLS 1.2", server_scan_result.scan_result.tls_1_2_cipher_suites),
            ("TLS 1.1", server_scan_result.scan_result.tls_1_1_cipher_suites),
            ("TLS 1.0", server_scan_result.scan_result.tls_1_0_cipher_suites),
            ("SSL 3.0", server_scan_result.scan_result.ssl_3_0_cipher_suites),
            ("SSL 2.0", server_scan_result.scan_result.ssl_2_0_cipher_suites),
        ]:
            if attempt.status == ScanCommandAttemptStatusEnum.COMPLETED:
                print(f"\nAccepted cipher suites for {version}:")
                safe, unsafe = mark_and_sort_cipher_suites(attempt.result.accepted_cipher_suites)
                result_entry["ciphers"][version] = [(suite, "SAFE") for suite in safe] + [
                    (suite, "UNSAFE") for suite in unsafe
                ]

        additional_checks = {
            "Heartbleed": server_scan_result.scan_result.heartbleed,
            "Session Resumption": server_scan_result.scan_result.session_resumption,
            "TLS Compression": server_scan_result.scan_result.tls_compression,
            "TLS Fallback SCSV": server_scan_result.scan_result.tls_fallback_scsv,
            "ROBOT": server_scan_result.scan_result.robot,
            "Session Renegotiation": server_scan_result.scan_result.session_renegotiation,
            "Elliptic Curves": server_scan_result.scan_result.elliptic_curves,
            "HTTP Headers": server_scan_result.scan_result.http_headers,
        }

        for check_name, attempt in additional_checks.items():
            if attempt.status == ScanCommandAttemptStatusEnum.COMPLETED:
                result_entry["additional_checks"].append((check_name, "SUCCESS"))
                print(f"{check_name}: {Fore.GREEN}SUCCESS{Style.RESET_ALL}")
            else:
                result_entry["additional_checks"].append((check_name, "FAILED"))
                print(f"{check_name}: {Fore.RED}FAILED{Style.RESET_ALL}")

        output_file = f"{hostname}_ssl_scan_results.docx"
        write_results_to_docx(result_entry, output_file)

if __name__ == "__main__":
    main()
