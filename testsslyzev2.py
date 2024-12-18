from datetime import datetime, timezone
from pathlib import Path
from typing import List
from colorama import init, Fore, Style
from docx import Document
from docx.shared import Pt
from sslyze import (
    Scanner,
    ServerScanRequest,
    ScanCommand,
    ServerNetworkLocation,
    ServerScanStatusEnum,
)
import warnings
from cryptography.utils import CryptographyDeprecationWarning

# Suppress cryptography deprecation warnings
warnings.filterwarnings("ignore", category=CryptographyDeprecationWarning)

# Initialize colorama for colored terminal output
init(autoreset=True)

# Define Mozilla's intermediate TLS cipher suites for each TLS version
MOZILLA_INTERMEDIATE_CIPHERS = {
    "TLS 1.3": {
        "TLS_AES_128_GCM_SHA256",
        "TLS_AES_256_GCM_SHA384",
        "TLS_CHACHA20_POLY1305_SHA256",
    },
    "TLS 1.2": {
        "TLS_ECDHE_ECDSA_WITH_AES_128_GCM_SHA256",
        "TLS_ECDHE_RSA_WITH_AES_128_GCM_SHA256",
        "TLS_ECDHE_ECDSA_WITH_AES_256_GCM_SHA384",
        "TLS_ECDHE_RSA_WITH_AES_256_GCM_SHA384",
        "TLS_ECDHE_ECDSA_WITH_CHACHA20_POLY1305_SHA256",
        "TLS_ECDHE_RSA_WITH_CHACHA20_POLY1305_SHA256",
        "TLS_DHE_RSA_WITH_AES_128_GCM_SHA256",
        "TLS_DHE_RSA_WITH_AES_256_GCM_SHA384",
        "TLS_DHE_RSA_WITH_CHACHA20_POLY1305_SHA256",
    },
}

DEPRECATED_SSL_VERSIONS = {"SSL 2.0", "SSL 3.0"}

ADDITIONAL_CHECKS = {
    "ELLIPTIC_CURVES": ScanCommand.ELLIPTIC_CURVES,
    "ROBOT": ScanCommand.ROBOT,
    "SESSION_RESUMPTION": ScanCommand.SESSION_RESUMPTION,
    "TLS_COMPRESSION": ScanCommand.TLS_COMPRESSION,
    "TLS_1_3_EARLY_DATA": ScanCommand.TLS_1_3_EARLY_DATA,
    "TLS_FALLBACK_SCSV": ScanCommand.TLS_FALLBACK_SCSV,
    "HEARTBLEED": ScanCommand.HEARTBLEED,
    "HTTP_HEADERS": ScanCommand.HTTP_HEADERS,
    "OPENSSL_CCS_INJECTION": ScanCommand.OPENSSL_CCS_INJECTION,
    "SESSION_RENEGOTIATION": ScanCommand.SESSION_RENEGOTIATION,
}

def print_ascii_banner():
    banner = r"""
    ==========================================
      ____           _         _   ____  _    
     / ___|   _ __  | |   ___ | |_| __ )| |   
     \___ \  | '_ \ | |  / _ \| __|  _ \| |   
      ___) | | |_) || |_|  __/| |_| |_) | |___
     |____/  | .__/ |____\___| \__|____/|_____|
             |_|                              
    SSL/TLS Scan with Mozilla Compliance Check
    ==========================================
    """
    print(Fore.CYAN + banner + Style.RESET_ALL)

def mark_and_sort_cipher_suites(cipher_suites, tls_version):
    safe_suites = []
    unsafe_suites = []

    for suite in cipher_suites:
        cipher_name = suite.cipher_suite.name
        if cipher_name in MOZILLA_INTERMEDIATE_CIPHERS.get(tls_version, set()):
            safe_suites.append(cipher_name)
        else:
            unsafe_suites.append(cipher_name)

    for suite in sorted(safe_suites):
        print(f"{Fore.GREEN}* {suite} - SAFE{Style.RESET_ALL}")

    for suite in sorted(unsafe_suites):
        print(f"{Fore.RED}* {suite} - UNSAFE{Style.RESET_ALL}")

    return sorted(safe_suites), sorted(unsafe_suites)

def write_results_to_docx(results: dict, output_file: str):
    doc = Document()
    doc.add_heading(f"SSL/TLS Scan Results for {results['hostname']}", level=1)

    for tls_version, ciphers in results["ciphers"].items():
        doc.add_heading(f"{tls_version} Cipher Suites", level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Cipher Suite"
        hdr_cells[1].text = "Status"
        for cipher, status in ciphers:
            row_cells = table.add_row().cells
            row_cells[0].text = cipher
            row_cells[1].text = status

    doc.add_heading("Additional Checks", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Check"
    hdr_cells[1].text = "Result"
    hdr_cells[2].text = "Details"
    for check_name, result, details in results["additional_checks"]:
        row_cells = table.add_row().cells
        row_cells[0].text = check_name
        row_cells[1].text = result
        row_cells[2].text = details

    doc.save(output_file)
    print(f"\nResults saved to {output_file}")

def main() -> None:
    print_ascii_banner()
    print("=> Starting the scans")

    servers_input = input("Enter server hostnames separated by commas: ")
    server_list = [host.strip() for host in servers_input.split(",") if host.strip()]

    if not server_list:
        print("No servers provided. Exiting.")
        return

    all_scan_requests = []
    for server in server_list:
        try:
            all_scan_requests.append(ServerScanRequest(
                server_location=ServerNetworkLocation(hostname=server),
                scan_commands=[
                    ScanCommand.SSL_2_0_CIPHER_SUITES,
                    ScanCommand.SSL_3_0_CIPHER_SUITES,
                    ScanCommand.TLS_1_0_CIPHER_SUITES,
                    ScanCommand.TLS_1_1_CIPHER_SUITES,
                    ScanCommand.TLS_1_2_CIPHER_SUITES,
                    ScanCommand.TLS_1_3_CIPHER_SUITES,
                    *ADDITIONAL_CHECKS.values(),
                ]
            ))
        except Exception as e:
            print(f"Error: Could not prepare scan for hostname '{server}': {e}")
            continue

    scanner = Scanner()
    scanner.queue_scans(all_scan_requests)

    for server_scan_result in scanner.get_results():
        hostname = server_scan_result.server_location.hostname
        print(f"\n\n****Results for {hostname}****")

        if server_scan_result.scan_status == ServerScanStatusEnum.ERROR_NO_CONNECTIVITY:
            print(f"Error: Could not connect to {hostname}: {server_scan_result.connectivity_error_trace}")
            continue

        result_entry = {"hostname": hostname, "ciphers": {}, "additional_checks": []}

        for version, command in [
            ("SSL 2.0", server_scan_result.scan_result.ssl_2_0_cipher_suites),
            ("SSL 3.0", server_scan_result.scan_result.ssl_3_0_cipher_suites),
            ("TLS 1.0", server_scan_result.scan_result.tls_1_0_cipher_suites),
            ("TLS 1.1", server_scan_result.scan_result.tls_1_1_cipher_suites),
            ("TLS 1.2", server_scan_result.scan_result.tls_1_2_cipher_suites),
            ("TLS 1.3", server_scan_result.scan_result.tls_1_3_cipher_suites),
        ]:
            if version in DEPRECATED_SSL_VERSIONS:
                if command and hasattr(command.result, "accepted_cipher_suites") and command.result.accepted_cipher_suites:
                    print(f"{Fore.RED}{version} detected! This is a critical failure. It must be disabled.{Style.RESET_ALL}")
                    result_entry["ciphers"][version] = [(suite.cipher_suite.name, "CRITICAL FAILURE") for suite in command.result.accepted_cipher_suites]
                else:
                    print(f"{Fore.GREEN}{version} not detected. Server is safe from this deprecated protocol.{Style.RESET_ALL}")
                    result_entry["ciphers"][version] = [("NONE", "NOT DETECTED")]
                continue

            if command:
                if hasattr(command.result, "accepted_cipher_suites"):
                    print(f"\nAccepted cipher suites for {version}:")
                    safe, unsafe = mark_and_sort_cipher_suites(command.result.accepted_cipher_suites, version)
                    result_entry["ciphers"][version] = [(suite, "SAFE") for suite in safe] + [
                        (suite, "UNSAFE") for suite in unsafe
                    ]
                else:
                    print(f"{Fore.YELLOW}No cipher suites available for {version}.{Style.RESET_ALL}")
                    result_entry["ciphers"][version] = []

        for check_name, command in ADDITIONAL_CHECKS.items():
            check_result = getattr(server_scan_result.scan_result, command.name.lower(), None)
            if check_result:
                if hasattr(check_result, "is_vulnerable") and check_result.is_vulnerable:
                    result_entry["additional_checks"].append((check_name, "FAILED", "Vulnerable"))
                    print(f"{check_name}: {Fore.RED}FAILED{Style.RESET_ALL} - Vulnerable")
                else:
                    result_entry["additional_checks"].append((check_name, "SUCCESS", "Not Vulnerable"))
                    print(f"{check_name}: {Fore.GREEN}SUCCESS{Style.RESET_ALL} - Not Vulnerable")
            else:
                result_entry["additional_checks"].append((check_name, "FAILED", "Not Implemented"))
                print(f"{check_name}: {Fore.RED}FAILED{Style.RESET_ALL} - Not Implemented")

        output_file = f"{hostname}_ssl_scan_results.docx"
        write_results_to_docx(result_entry, output_file)

if __name__ == "__main__":
    main()
